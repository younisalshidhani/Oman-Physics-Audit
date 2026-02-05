# app.py — نسخة كاملة معدّلة للتحكم في اختيار النموذج ومعالجة الأخطاء بشكل أفضل
# ملاحظات: انسخ هذا الملف كاملاً واصله محل app.py في مستودعك أو على جهازك.
import streamlit as st
import fitz  # PyMuPDF
import google.generativeai as genai
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import traceback

st.set_page_config(page_title="نظام تدقيق الاختبارات العماني", layout="wide")

st.markdown("""
    <style>
    .stApp { direction: rtl; text-align: right; }
    div[data-testid="stSidebar"] { text-align: right; direction: rtl; }
    div[data-testid="stMarkdownContainer"] { text-align: right; direction: rtl; }
    table { width: 100%; border-collapse: collapse; direction: rtl; margin-bottom: 20px; }
    th, td { border: 1px solid #ddd; padding: 10px; text-align: right; }
    th { background-color: #f8f9fa; }
    .error-box { background:#fde2e2; padding:12px; border-radius:6px; color:#900; }
    </style>
    """, unsafe_allow_html=True)

# ----------------------------
# دوال مساعدة
# ----------------------------
def extract_pdf_text(uploaded_file):
    """يقرأ ملف UploadedFile من Streamlit ويُرجع نص كل صفحاته"""
    if not uploaded_file:
        return ""
    data = uploaded_file.read()
    if not data:
        return ""
    try:
        doc = fitz.open(stream=data, filetype="pdf")
        texts = []
        for page in doc:
            texts.append(page.get_text())
        return "\n".join(texts)
    except Exception as e:
        return f"[خطأ في استخراج النص من PDF: {e}]"

def generate_word(data, subject, grade, semester, exam_type):
    doc = Document()
    header = doc.add_heading(f'تقرير تحليل اختبار {subject}', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para = doc.add_paragraph(f'الصف: {grade} | الفصل: {semester} | النوع: {exam_type}')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("-" * 50)

    # جدول 1 - تحليل المفردات
    doc.add_heading('1. تحليل المفردات', level=1)
    table1 = doc.add_table(rows=1, cols=6)
    table1.style = 'Table Grid'
    hdrs = ["م", "الهدف", "المستوى", "الدرجة", "الملاحظة", "التعديل"]
    for i, h in enumerate(hdrs):
        table1.rows[0].cells[i].text = h
    for item in data.get("vocab", []):
        row = table1.add_row().cells
        row[0].text = str(item.get("q", ""))
        row[1].text = item.get("obj", "")
        row[2].text = item.get("level", "")
        row[3].text = str(item.get("mark", ""))
        row[4].text = item.get("note", "")
        row[5].text = item.get("fix", "")

    # جدول 2 - المواصفات
    doc.add_heading('2. الجدول العامل', level=1)
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdrs2 = ["البند", "البيان", "التقييم"]
    for i, h in enumerate(hdrs2):
        table2.rows[0].cells[i].text = h
    specs = data.get("specs", {})
    mapping = {"q_count":"عدد الأسئلة", "lessons":"تغطية الدروس", "ao1":"درجات AO1", "ao2":"درجات AO2", "clarity":"الوضوح"}
    for key, label in mapping.items():
        row = table2.add_row().cells
        item = specs.get(key, {})
        row[0].text = label
        row[1].text = str(item.get("val", "-"))
        row[2].text = item.get("status", "-")

    doc.add_heading('3. التقدير العام', level=1)
    doc.add_paragraph(data.get("summary", ""))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def choose_model_and_list():
    """
    يحاول جلب قائمة النماذج المتاحة ثم يختار اسماً مُفضلاً إن وُجد.
    يُرجع tuple: (chosen_model_name or None, raw_list_repr)
    """
    try:
        # بعض إصدارات المكتبة تعرض واجهة مختلفة لاستخراج النماذج، نجرب أكثر من طريقة
        list_result = None
        if hasattr(genai, "list_models"):
            list_result = genai.list_models()
        elif hasattr(genai, "models") and hasattr(genai.models, "list"):
            list_result = genai.models.list()
        else:
            # لا توجد واجهة list models معروفة - نُعيد None
            return None, "list_models not supported in this client version."

        # نُحضّر تمثيلاً نصياً للنتيجة لعرضه في حال الخطأ
        try:
            raw_list = json.dumps(list_result, default=lambda o: o.__dict__, ensure_ascii=False)
        except Exception:
            raw_list = str(list_result)

        # نحاول استخراج أسماء النماذج إن وُجدت في هيكلية معروفة
        names = []
        if isinstance(list_result, dict):
            # بعض الواجهات تعيد dict مع مفتاح 'models'
            for k in ("models", "model_versions", "data"):
                if k in list_result and isinstance(list_result[k], list):
                    for entry in list_result[k]:
                        if isinstance(entry, dict) and entry.get("name"):
                            names.append(entry.get("name"))
                        elif hasattr(entry, "name"):
                            names.append(getattr(entry, "name"))
        elif isinstance(list_result, list):
            for entry in list_result:
                if isinstance(entry, dict) and entry.get("name"):
                    names.append(entry.get("name"))
                elif hasattr(entry, "name"):
                    names.append(getattr(entry, "name"))
        else:
            # fallback: stringify
            # try to parse any "name": "..." patterns
            raw = str(list_result)
            # بسيط: البحث عن كلمات models/... في النص
            tokens = raw.split()
            for t in tokens:
                if "models/" in t or "gemini" in t or "bison" in t:
                    cleaned = t.strip(",\"'")
                    names.append(cleaned)

        # ترتيب التفضيل لاختيار نموذج موثوق إن وُجد
        preferred = ["models/gemini-1.5-flash", "models/gemini-1.5", "models/gemini-1.0", "models/text-bison-001", "bison", "gemini"]
        for p in preferred:
            for n in names:
                if p in n:
                    return n, raw_list

        # إن لم نجد تفضيلات، نُعيد النموذج الأول إن وُجد
        if names:
            return names[0], raw_list

        return None, raw_list
    except Exception as e:
        return None, f"Error calling list_models: {e}"

# ----------------------------
# واجهة جانبية
# ----------------------------
with st.sidebar:
    st.header("⚙️ الإعدادات")
    api_key = st.text_input("مفتاح API (لن يُخزن):", type="password")
    subject = st.selectbox("المادة:", ["فيزياء", "كيمياء", "أحياء", "علوم"])
    grade = st.selectbox("الصف:", ["11", "12"])
    semester = st.selectbox("الفصل الدراسي:", ["الأول", "الثاني"])
    exam_type = st.selectbox("نوع الاختبار:", ["قصير", "تجريبي/نهائي"])
    pages = st.text_input("نطاق الصفحات (مثلاً 1-5):", "1-50")

# ----------------------------
# الجسم الرئيسي
# ----------------------------
st.title(f"🔍 تدقيق اختبار {subject} - الصف {grade}")

c1, c2, c3 = st.columns(3)
with c1:
    f_test = st.file_uploader("ملف الاختبار (PDF)", type="pdf")
with c2:
    f_policy = st.file_uploader("وثيقة التقويم (PDF)", type="pdf")
with c3:
    f_book = st.file_uploader("كتاب الطالب (PDF)", type="pdf")

if st.button("🚀 تنفيذ التحليل"):
    if not api_key:
        st.error("يرجى إدخال مفتاح API في الإعدادات أولاً.")
    elif not f_test:
        st.error("يرجى رفع ملف الاختبار.")
    else:
        genai.configure(api_key=api_key)

        with st.spinner("جاري التحليل وفق المعايير..."):
            t_test = extract_pdf_text(f_test)
            t_policy = extract_pdf_text(f_policy) if f_policy else ""
            t_book = extract_pdf_text(f_book) if f_book else ""

            # لتجنب إرسال نصوص ضخمة جداً، نقتصر على مقتطفات أولية — يمكن تعديل max_chunk حسب الحاجة
            max_chunk = 4000
            t_test_snip = t_test[:max_chunk]
            t_policy_snip = t_policy[:max_chunk]
            t_book_snip = t_book[:max_chunk]

            prompt = f"""
حلل اختبار {subject} للصف {grade} فصل {semester}.
قارن أسئلة الاختبار مع وثيقة التقويم (أدخل البنود المطابقة إن وُجدت) وبمحتوى كتاب الطالب.
- وثيقة التقويم (مقتطف): {t_policy_snip}
- كتاب الطالب (مقتطف): صفحات {pages} => {t_book_snip}
- نص الاختبار (مقتطف): {t_test_snip}

اطرح مخرجات JSON فقط بالهيكل التالي:
{{
  "vocab":[
    {{
      "q": "رقم السؤال",
      "obj": "الهدف/المؤشر المتطابق",
      "level": "AO1|AO2|AO3|...",
      "mark": 1,
      "note": "ملاحظات",
      "fix": "اقتراح تعديل"
    }}
  ],
  "specs": {{
    "q_count":{{"val": 0, "status": "OK|Missing"}},
    "lessons":{{"val":"قائمة","status":"Covered|Not covered"}},
    "ao1":{{"val":0,"status":"OK"}},
    "ao2":{{"val":0,"status":"OK"}},
    "clarity":{{"val":"High|Low","status":"OK|Needs revision"}}
  }},
  "summary":"ملخص تحليلي قصير"
}}
"""

            # اختيار نموذج مناسب من قائمة النماذج المتاحة
            chosen_model, models_raw = choose_model_and_list()

            if not chosen_model:
                st.error("تعذر العثور على نموذج متاح عبر واجهة API المثبتة. انظر تفاصيل الخطأ أدناه.")
                st.code(models_raw)
            else:
                # نعرض اسم النموذج الذي سنحاول استخدامه لمساعدة التشخيص
                st.info(f"استخدام النموذج: {chosen_model}")

                raw_text = None
                gen_error = None
                try:
                    # محاولة استخدام أسلوب قديم إن أمكن (caching for compatibility)
                    try:
                        # بعض النسخ من المكتبة توفر مولد نموذج ككلاس
                        model_obj = None
                        if hasattr(genai, "GenerativeModel"):
                            try:
                                model_obj = genai.GenerativeModel(chosen_model)
                            except Exception:
                                model_obj = None
                        if model_obj is not None and hasattr(model_obj, "generate_content"):
                            res = model_obj.generate_content(prompt)
                            raw_text = getattr(res, "text", None) or getattr(res, "content", None) or str(res)
                        else:
                            raise AttributeError("model_obj.generate_content not available - falling back")
                    except Exception:
                        # محاولة واجهة genai.generate_text (شائعة في بعض إصدارات المكتبة)
                        try:
                            if hasattr(genai, "generate_text"):
                                res = genai.generate_text(model=chosen_model, prompt=prompt)
                                # شكل المحتوى قد يختلف حسب النسخة
                                raw_text = getattr(res, "text", None) or getattr(res, "content", None) or getattr(res, "output", None) or str(res)
                            else:
                                raise AttributeError("genai.generate_text not available")
                        except Exception:
                            # محاولة واجهة عامة genai.generate إن وُجدت
                            if hasattr(genai, "generate"):
                                res = genai.generate(model=chosen_model, prompt=prompt)
                                raw_text = getattr(res, "text", None) or getattr(res, "content", None) or getattr(res, "output", None) or str(res)
                            else:
                                raise RuntimeError("لا توجد واجهة مدعومة للتوليد في مكتبة google-generativeai المثبتة.")

                    # تنظيف وإيجاد JSON
                    if not raw_text:
                        raise RuntimeError("المولد أعاد نتيجة فارغة.")

                    cleaned = raw_text.replace("```json", "").replace("```", "").strip()
                    start = cleaned.find("{")
                    end = cleaned.rfind("}")
                    if start == -1 or end == -1:
                        raise ValueError("لم يتم العثور على JSON صالح في ناتج النموذج.")
                    js_str = cleaned[start:end+1]
                    data = json.loads(js_str)

                    st.success("اكتمل التحليل")

                    # عرض تحليل المفردات (جدول HTML بسيط)
                    st.subheader("تحليل المفردات")
                    vocab = data.get("vocab", [])
                    if vocab:
                        rows = "".join([f"<tr><td>{i.get('q','')}</td><td>{i.get('obj','')}</td><td>{i.get('level','')}</td><td>{i.get('mark','')}</td><td>{i.get('note','')}</td><td>{i.get('fix','')}</td></tr>" for i in vocab])
                        st.markdown(f"<table><tr><th>م</th><th>الهدف</th><th>المستوى</th><th>الدرجة</th><th>الملاحظة</th><th>التعديل</th></tr>{rows}</table>", unsafe_allow_html=True)
                    else:
                        st.info("لا توجد مفردات مُعالجة في النتيجة.")

                    # عرض الجدول العامل
                    st.subheader("الجدول العامل")
                    specs = data.get("specs", {})
                    mapping = {"q_count":"العدد","lessons":"الدروس","ao1":"AO1","ao2":"AO2","clarity":"الوضوح"}
                    s_rows = ""
                    for k, lbl in mapping.items():
                        val = specs.get(k, {}).get("val", "-")
                        status = specs.get(k, {}).get("status", "-")
                        s_rows += f"<tr><td>{lbl}</td><td>{val}</td><td>{status}</td></tr>"
                    st.markdown(f"<table><tr><th>البند</th><th>القيمة</th><th>الحالة</th></tr>{s_rows}</table>", unsafe_allow_html=True)

                    st.info(data.get("summary", ""))

                    st.download_button("📥 تحميل التقرير (Word)", generate_word(data, subject, grade, semester, exam_type), "Report.docx")

                except Exception as gen_exc:
                    # لا نعرض مفتاح الـAPI أبداً؛ نعرض الأخطاء التفصيلية لتشخيص المشكلة
                    gen_error = str(gen_exc)
                    st.error("فشل التحليل — راجع التفاصيل أدناه.")
                    st.markdown(f"<div class='error-box'>خطأ التنفيذ: {gen_error}</div>", unsafe_allow_html=True)
                    # نعرض مخرجات النموذج الخام إن وُجدت لمساعدة التشخيص
                    if raw_text:
                        st.subheader("مخرجات النموذج (خام)")
                        st.code(raw_text)
                    # نعرض قائمة النماذج الخام لمساعدتك على اختيار اسم نموذج يدعمه حسابك/المكتبة
                    st.subheader("قائمة النماذج (لمساعدة التشخيص)")
                    st.code(models_raw)

                    # احتياطي: لوج الاستثناء الكامل داخل اللوحة (للمطور فقط، لكن ليس مفتاح الـAPI)
                    st.subheader("تفاصيل الاستثناء (traceback)")
                    st.code(traceback.format_exc())
